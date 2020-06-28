import docx, os

# doc = docx.Document(r'C:\Users\Jai\Desktop\Docs\demo.docx')
os.chdir(r"C:\Users\Jai\Desktop\Docs")


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


"""
Test:
from PDFandWorddocs import readDocx
print(readDocx.getText('demo.docx'))
"""
